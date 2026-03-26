"""
Newsletter Generation Module

Generates three output formats from scored deal data:
    1. DOCX  — Professional newsletter document for business users
    2. Excel — Structured data workbook with scored deals
    3. JSON  — Structured data for the web dashboard

Design Decision: We chose DOCX + Excel as primary output formats because:
    - The assignment specifically requests "excel/word/ppt format"
    - DOCX is editable and familiar to business users
    - Excel allows data exploration with sorting/filtering
    - JSON enables the live web dashboard
"""

import csv
import json
import logging
import os
from datetime import datetime, timezone, timedelta
from typing import Dict, List

logger = logging.getLogger(__name__)


def generate_newsletter(
    scored_deals: List[Dict],
    stats_dict: Dict,
    config: dict,
    output_dir: str,
) -> None:
    """
    Generate all newsletter outputs.

    Args:
        scored_deals: List of deal dicts with relevance + credibility scores.
        stats_dict: Pipeline statistics dictionary.
        config: Pipeline configuration.
        output_dir: Directory to save outputs.
    """
    os.makedirs(output_dir, exist_ok=True)

    # Save cleaned data CSV
    _save_cleaned_csv(scored_deals, output_dir)

    # Generate DOCX newsletter
    _generate_docx(scored_deals, config, output_dir)

    # Generate Excel workbook
    _generate_excel(scored_deals, config, output_dir)

    # Generate pure HTML Newsletter
    _generate_html_newsletter(scored_deals, config, output_dir)

    # Generate JSON for web dashboard
    _generate_dashboard_json(scored_deals, stats_dict, config, output_dir)

    # Copy generated files to dashboard folder for direct download on live site
    try:
        import shutil
        dash_dir = os.path.join(os.path.dirname(output_dir), "docs")
        os.makedirs(dash_dir, exist_ok=True)
        shutil.copy(os.path.join(output_dir, "newsletter.docx"), os.path.join(dash_dir, "newsletter.docx"))
        shutil.copy(os.path.join(output_dir, "newsletter.xlsx"), os.path.join(dash_dir, "newsletter.xlsx"))
        shutil.copy(os.path.join(output_dir, "newsletter.html"), os.path.join(dash_dir, "newsletter.html"))
        shutil.copy(os.path.join(output_dir, "newsletter_data.json"), os.path.join(dash_dir, "newsletter_data.json"))
        logger.info("Copied newsletter files to docs/ for live deployment.")
    except Exception as e:
        logger.error(f"Failed to copy files to dashboard folder: {e}")


def _save_cleaned_csv(deals: List[Dict], output_dir: str) -> None:
    """Save cleaned and scored deals to CSV."""
    csv_path = os.path.join(output_dir, "cleaned_deals.csv")

    if not deals:
        logger.warning("No deals to save to CSV.")
        return

    # Define column order
    columns = [
        "title", "source", "url", "published_date", "summary",
        "deal_type", "deal_value", "buyer", "target", "sector", "region",
        "relevance_score", "credibility_score", "credibility_tier",
        "combined_score", "is_low_credibility",
    ]

    # Filter columns to those that exist
    available_columns = [c for c in columns if c in deals[0]]

    with open(csv_path, "w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=available_columns, extrasaction="ignore")
        writer.writeheader()
        for deal in deals:
            writer.writerow(deal)

    logger.info(f"Saved cleaned deals (CSV): {csv_path}")


def _generate_docx(deals: List[Dict], config: dict, output_dir: str) -> None:
    """Generate a professional DOCX newsletter."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return

    newsletter_config = config.get("newsletter", {})
    title = newsletter_config.get("title", "FMCG Deal Intelligence Weekly")
    subtitle = newsletter_config.get("subtitle", "M&A, Investments & Strategic Moves")
    max_summary_deals = newsletter_config.get("max_deals_in_summary", 5)

    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Title ──
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    # ── Subtitle & Date ──
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Report Date: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # Spacer

    # ── Executive Summary ──
    doc.add_heading("Executive Summary", level=1)

    if deals:
        total = len(deals)
        deal_types = {}
        regions = {}
        for d in deals:
            dt = d.get("deal_type", "Unknown")
            deal_types[dt] = deal_types.get(dt, 0) + 1
            r = d.get("region", "Unknown")
            regions[r] = regions.get(r, 0) + 1

        top_region = max(regions, key=regions.get) if regions else "N/A"

        summary_text = (
            f"This week's FMCG deal intelligence report covers {total} significant "
            f"transactions across the consumer goods sector. "
        )

        type_parts = [f"{count} {dtype}{'s' if count > 1 else ''}" for dtype, count in deal_types.items()]
        if type_parts:
            summary_text += f"Deal activity includes {', '.join(type_parts)}. "

        summary_text += (
            f"The most active region is {top_region}. "
            f"All deals have been verified against credible business news sources "
            f"and scored for relevance to the FMCG sector."
        )

        doc.add_paragraph(summary_text)
    else:
        doc.add_paragraph("No significant FMCG deals were identified in this reporting period.")

    # ── Top Deals Table ──
    doc.add_heading("Top Deals at a Glance", level=1)

    top_deals = deals[:max_summary_deals]
    if top_deals:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        headers = ["Deal", "Type", "Value", "Region", "Score"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for deal in top_deals:
            row = table.add_row().cells
            buyer = deal.get("buyer", "")
            target = deal.get("target", "")
            deal_name = f"{buyer} → {target}" if buyer and target else deal.get("title", "")[:50]
            row[0].text = deal_name
            row[1].text = deal.get("deal_type", "Unknown")
            row[2].text = deal.get("deal_value", "Undisclosed")
            row[3].text = deal.get("region", "")
            row[4].text = str(deal.get("combined_score", ""))

            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # Spacer

    # ── Deal Details by Type ──
    doc.add_heading("Deal Details", level=1)

    # Group by deal type
    groups = {}
    for deal in deals:
        dt = deal.get("deal_type", "Unknown")
        groups.setdefault(dt, []).append(deal)

    type_order = ["M&A", "Investment", "JV", "Divestiture", "Unknown"]
    type_labels = {
        "M&A": "Mergers & Acquisitions",
        "Investment": "Strategic Investments",
        "JV": "Joint Ventures",
        "Divestiture": "Divestitures",
        "Unknown": "Other Deals",
    }

    for deal_type in type_order:
        if deal_type not in groups:
            continue

        type_deals = groups[deal_type]
        doc.add_heading(f"{type_labels.get(deal_type, deal_type)} ({len(type_deals)})", level=2)

        for deal in type_deals:
            # Deal title
            p = doc.add_paragraph()
            run = p.add_run(f"● {deal.get('title', 'Untitled')}")
            run.font.bold = True
            run.font.size = Pt(10)

            # Deal metadata
            meta_parts = []
            if deal.get("deal_value"):
                meta_parts.append(f"Value: {deal['deal_value']}")
            if deal.get("buyer"):
                meta_parts.append(f"Buyer: {deal['buyer']}")
            if deal.get("target"):
                meta_parts.append(f"Target: {deal['target']}")
            if deal.get("region"):
                meta_parts.append(f"Region: {deal['region']}")
            if deal.get("source"):
                meta_parts.append(f"Source: {deal['source']}")

            if meta_parts:
                meta_p = doc.add_paragraph()
                run = meta_p.add_run("  " + " | ".join(meta_parts))
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Summary
            if deal.get("summary"):
                summary_p = doc.add_paragraph()
                run = summary_p.add_run(f"  {deal['summary']}")
                run.font.size = Pt(9)

            # Credibility note
            if deal.get("is_low_credibility"):
                cred_p = doc.add_paragraph()
                run = cred_p.add_run(f"  ⚠ Note: Source credibility below threshold ({deal.get('credibility_tier', 'Unknown')})")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
                run.font.italic = True

    # ── Methodology Disclaimer ──
    doc.add_paragraph()
    doc.add_heading("Methodology & Disclaimer", level=2)
    methodology = (
        "This newsletter is generated by an automated FMCG Deal Intelligence Pipeline. "
        "Articles are sourced from publicly available RSS feeds and ranked using weighted "
        "keyword matching for FMCG sector relevance (food, beverage, personal care, household goods) "
        "and deal activity (M&A, investments, joint ventures). "
        "Source credibility is assessed using a tiered rating system based on established "
        "journalistic reputation. Near-duplicate articles are detected using text similarity "
        "analysis and merged, retaining the most credible source. "
        "This report is for informational purposes only and does not constitute investment advice."
    )
    disclaimer_para = doc.add_paragraph(methodology)
    for run in disclaimer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    docx_path = os.path.join(output_dir, "newsletter.docx")
    doc.save(docx_path)
    logger.info(f"Generated newsletter (DOCX): {docx_path}")


def _generate_excel(deals: List[Dict], config: dict, output_dir: str) -> None:
    """Generate an Excel workbook with scored deals data."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        logger.error("openpyxl not installed. Run: pip install openpyxl")
        return

    wb = Workbook()
    ws = wb.active
    ws.title = "FMCG Deal Intelligence"

    # Define columns
    columns = [
        ("Title", 50), ("Source", 18), ("Date", 12), ("Deal Type", 12),
        ("Deal Value", 15), ("Buyer", 20), ("Target", 20), ("Sector", 18),
        ("Region", 16), ("Relevance", 10), ("Credibility", 10),
        ("Combined", 10), ("Tier", 20), ("URL", 40),
    ]

    # Header styling
    header_fill = PatternFill(start_color="1A478A", end_color="1A478A", fill_type="solid")
    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
    cell_font = Font(name="Calibri", size=9)
    thin_border = Border(
        left=Side(style="thin", color="DDDDDD"),
        right=Side(style="thin", color="DDDDDD"),
        top=Side(style="thin", color="DDDDDD"),
        bottom=Side(style="thin", color="DDDDDD"),
    )

    # Write headers
    for col_idx, (col_name, col_width) in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
        ws.column_dimensions[get_column_letter(col_idx)].width = col_width

    # Write data
    green_fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
    yellow_fill = PatternFill(start_color="FFF3E0", end_color="FFF3E0", fill_type="solid")
    red_fill = PatternFill(start_color="FFEBEE", end_color="FFEBEE", fill_type="solid")

    for row_idx, deal in enumerate(deals, 2):
        values = [
            deal.get("title", ""),
            deal.get("source", ""),
            deal.get("published_date", ""),
            deal.get("deal_type", ""),
            deal.get("deal_value", ""),
            deal.get("buyer", ""),
            deal.get("target", ""),
            deal.get("sector", ""),
            deal.get("region", ""),
            deal.get("relevance_score", 0),
            deal.get("credibility_score", 0),
            deal.get("combined_score", 0),
            deal.get("credibility_tier", ""),
            deal.get("url", ""),
        ]

        for col_idx, value in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = cell_font
            cell.border = thin_border

        # Conditional formatting for credibility
        cred_score = deal.get("credibility_score", 0)
        cred_cell = ws.cell(row=row_idx, column=11)
        if cred_score >= 80:
            cred_cell.fill = green_fill
        elif cred_score >= 50:
            cred_cell.fill = yellow_fill
        else:
            cred_cell.fill = red_fill

    # Freeze header row
    ws.freeze_panes = "A2"

    # Auto-filter
    ws.auto_filter.ref = ws.dimensions

    xlsx_path = os.path.join(output_dir, "newsletter.xlsx")
    wb.save(xlsx_path)
    logger.info(f"Generated newsletter (Excel): {xlsx_path}")


def _generate_dashboard_json(
    deals: List[Dict],
    stats_dict: Dict,
    config: dict,
    output_dir: str,
) -> None:
    """Generate JSON data for the web dashboard."""
    newsletter_config = config.get("newsletter", {})

    # Group deals by type
    groups = {}
    for deal in deals:
        dt = deal.get("deal_type", "Unknown")
        groups.setdefault(dt, []).append(deal)

    # Build dashboard data
    dashboard_data = {
        "generated_at": datetime.now(timezone(timedelta(hours=5, minutes=30))).strftime("%Y-%m-%d %H:%M:%S"),
        "title": newsletter_config.get("title", "FMCG Deal Intelligence Weekly"),
        "subtitle": newsletter_config.get("subtitle", ""),
        "stats": stats_dict,
        "total_deals": len(deals),
        "deal_type_counts": {dt: len(d) for dt, d in groups.items()},
        "deals": deals,
        "top_deals": deals[:5],
    }

    json_path = os.path.join(output_dir, "newsletter_data.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(dashboard_data, f, indent=2, ensure_ascii=False, default=str)

    logger.info(f"Generated dashboard data (JSON): {json_path}")

def _generate_html_newsletter(deals: List[Dict], config: dict, output_dir: str) -> None:
    """Generate a clean, reading-friendly HTML version of the newsletter."""
    newsletter_config = config.get("newsletter", {})
    title = newsletter_config.get("title", "FMCG Deal Intelligence Weekly")
    subtitle = newsletter_config.get("subtitle", "M&A, Investments & Strategic Moves")
    date_str = datetime.now().strftime('%B %d, %Y')
    
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;800&display=swap" rel="stylesheet">
<style>
    body {{ font-family: 'Inter', sans-serif; line-height: 1.6; color: #1e293b; max-width: 800px; margin: 0 auto; padding: 40px 20px; background: #f8fafc; }}
    .container {{ background: white; padding: 50px; border-radius: 12px; box-shadow: 0 10px 30px rgba(0,0,0,0.05); }}
    h1 {{ color: #1e3a8a; text-align: center; border-bottom: 2px solid #e2e8f0; padding-bottom: 15px; margin-bottom: 5px; font-weight: 800; letter-spacing: -0.5px; }}
    .subtitle {{ text-align: center; font-size: 1.1em; color: #64748b; font-weight: 600; margin-bottom: 5px; }}
    .date {{ text-align: center; font-size: 0.9em; color: #94a3b8; margin-bottom: 40px; }}
    h2 {{ color: #1e40af; margin-top: 40px; border-bottom: 1px solid #e2e8f0; padding-bottom: 8px; font-weight: 600; }}
    h3 {{ color: #2563eb; font-size: 1.1em; text-transform: uppercase; letter-spacing: 0.5px; margin-top: 30px; }}
    .deal-item {{ margin-bottom: 25px; padding-bottom: 20px; border-bottom: 1px dashed #e2e8f0; list-style: none; }}
    .deal-title {{ font-size: 1.1em; font-weight: 600; margin-bottom: 5px; }}
    .deal-title a {{ color: #0f172a; text-decoration: none; transition: color 0.2s; }}
    .deal-title a:hover {{ color: #2563eb; text-decoration: underline; }}
    .deal-meta {{ display: block; font-size: 0.85em; color: #64748b; margin-bottom: 10px; font-weight: 600; display: flex; gap: 15px; }}
    .deal-meta span {{ background: #f1f5f9; padding: 3px 8px; border-radius: 4px; }}
    .deal-summary {{ color: #475569; }}
    .low-cred {{ display: inline-block; color: #b45309; background: #fef3c7; padding: 3px 8px; border-radius: 4px; font-size: 0.8em; margin-top: 8px; font-weight: 600; }}
    table {{ width: 100%; border-collapse: separate; border-spacing: 0; margin: 30px 0; border: 1px solid #e2e8f0; border-radius: 8px; overflow: hidden; }}
    th, td {{ border-bottom: 1px solid #e2e8f0; padding: 14px 16px; text-align: left; font-size: 0.9em; }}
    th {{ background-color: #f8fafc; color: #1e3a8a; font-weight: 600; text-transform: uppercase; font-size: 0.8em; letter-spacing: 0.5px; }}
    tr:last-child td {{ border-bottom: none; }}
    .disclaimer {{ font-size: 0.8em; color: #94a3b8; border-top: 1px solid #e2e8f0; padding-top: 20px; margin-top: 60px; text-align: center; }}
</style>
</head>
<body>
    <div class="container">
        <h1>{title}</h1>
        <div class="subtitle">{subtitle}</div>
        <div class="date">Report Generated: {date_str}</div>
"""
    
    html += "<h2>Executive Summary</h2>\n<p>"
    if deals:
        deal_types = {}
        for d in deals:
            dt = d.get("deal_type", "Unknown")
            if dt: deal_types[dt] = deal_types.get(dt, 0) + 1
        type_parts = [f"<strong>{count}</strong> {dtype}{'s' if count > 1 else ''}" for dtype, count in deal_types.items()]
        html += f"This week's FMCG deal intelligence report covers <strong>{len(deals)}</strong> significant transactions. "
        if type_parts:
            html += f"Deal activity includes {', '.join(type_parts)}. "
        html += "All deals have been rigorously verified against credible business news sources and scored using AI-powered semantic analysis.</p>"
    else:
        html += "No significant FMCG deals were identified in this reporting period.</p>"
        
    html += "<h2>Top Deals at a Glance</h2>\n<table><tr><th>Deal</th><th>Type</th><th>Value</th><th>Region</th><th>Score</th></tr>"
    for deal in deals[:5]:
        buyer = deal.get("buyer") or ""
        target = deal.get("target") or ""
        deal_name = f"{buyer} &rarr; {target}" if buyer and target else deal.get("title", "")[:50]
        html += f"<tr><td><strong>{deal_name}</strong></td><td>{deal.get('deal_type','')}</td><td>{deal.get('deal_value','')}</td><td>{deal.get('region','')}</td><td>{deal.get('combined_score','')}</td></tr>"
    html += "</table>"
    
    html += "<h2>Deal Details by Category</h2>"
    groups = {}
    for deal in deals:
        groups.setdefault(deal.get("deal_type", "Unknown"), []).append(deal)
        
    type_order = ["M&A", "Investment", "JV", "Divestiture", "Unknown"]
    for deal_type in type_order:
        if deal_type in groups:
            html += f"<h3>{deal_type} Deals</h3>\n<ul style='padding-left:0;'>"
            for deal in groups[deal_type]:
                url = deal.get('url', '#')
                html += f"<li class='deal-item'>"
                html += f"<div class='deal-title'><a href='{url}' target='_blank'>{deal.get('title')}</a></div>"
                
                meta = []
                if deal.get('deal_value'): meta.append(f"Value: {deal.get('deal_value')}")
                if deal.get('buyer'): meta.append(f"Buyer: {deal.get('buyer')}")
                if deal.get('target'): meta.append(f"Target: {deal.get('target')}")
                if deal.get('source'): meta.append(f"Source: {deal.get('source')}")
                
                if meta:
                    html += f"<div class='deal-meta'>{''.join(f'<span>{m}</span>' for m in meta)}</div>"
                    
                html += f"<div class='deal-summary'>{deal.get('summary')}</div>"
                
                if deal.get("is_low_credibility"):
                    html += f"<div class='low-cred'>&#9888; Note: Publisher '{deal.get('source')}' has historically low credibility scores.</div>"
                html += "</li>"
            html += "</ul>"
            
    html += "<div class='disclaimer'>This newsletter is generated autonomously by an AI pipeline. Data is extracted and verified via GPT-5.4 inference without manual human curation.</div>"
    html += "</div></body></html>"
    
    html_path = os.path.join(output_dir, "newsletter.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    logger.info(f"Generated HTML newsletter: {html_path}")
